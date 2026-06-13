"""
generate_student_assignments.py
================================
Run this script to generate 5 realistic student Word assignment files.
Run with:  python generate_student_assignments.py
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

def add_header(doc, student_name, roll_no, subject, assignment_no, date):
    heading = doc.add_heading("Assignment Submission", 0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")
    info_table = doc.add_table(rows=5, cols=2)
    info_table.style = "Table Grid"
    labels = ["Student Name", "Roll Number", "Subject", "Assignment No.", "Date of Submission"]
    values = [student_name, roll_no, subject, assignment_no, date]
    for i, (label, value) in enumerate(zip(labels, values)):
        info_table.rows[i].cells[0].text = label
        info_table.rows[i].cells[1].text = value
    doc.add_paragraph("")
    doc.add_paragraph("-" * 60)
    doc.add_paragraph("")

def add_body(doc, content_blocks):
    for heading_text, body_text in content_blocks:
        if heading_text:
            doc.add_heading(heading_text, level=2)
        para = doc.add_paragraph(body_text)
        para.paragraph_format.space_after = Pt(10)

# STUDENT 1: Ravi Sharma — GENUINE ORIGINAL (Expected: CLEAN)
def create_ravi():
    doc = Document()
    add_header(doc, "Ravi Sharma", "10A-021", "Environmental Science",
               "Assignment 3 - Effects of Air Pollution on Human Health", "9 June 2026")
    add_body(doc, [
        ("Introduction", (
            "I live near the highway in Pune and honestly the air quality there is really bad. "
            "Some mornings there is this visible haze outside our window and my grandmother always "
            "keeps her inhaler ready before she steps out. That is actually what made me interested "
            "in this topic. Air pollution is not just a news headline for us, it is something we "
            "deal with every single day."
        )),
        ("What is Air Pollution", (
            "Air pollution basically means when harmful substances get mixed into the air we breathe. "
            "These substances can be gases like carbon monoxide and sulfur dioxide, or they can be "
            "tiny solid particles called particulate matter. The main sources are vehicles, factories, "
            "burning of crop stubble, and construction dust. In cities like Delhi and Mumbai the "
            "problem is much worse because of the high number of vehicles."
        )),
        ("How it Affects the Human Body", (
            "The most direct effect is on the lungs. When we breathe polluted air the particles go "
            "deep into our lungs and can cause diseases like asthma, bronchitis, and even lung cancer "
            "if the exposure is over many years. I read that PM2.5 particles are the most dangerous "
            "because they are so small they can pass through the lung tissue directly into the "
            "bloodstream. Once in the blood they can reach the heart and brain.\n\n"
            "My grandmother has chronic bronchitis and the doctor specifically told her to avoid going "
            "out on days when the AQI is above 150. I check the AQI app now every morning for her. "
            "It made me realise this is not just a general health issue, it specifically targets older "
            "people and children whose lungs are weaker."
        )),
        ("Effects Beyond the Lungs", (
            "What surprised me during my research was that air pollution does not only affect "
            "breathing. Studies have found links between long term exposure to pollution and heart "
            "disease, stroke, and even mental health problems like depression and anxiety. Some "
            "research suggests that children growing up in heavily polluted areas score lower on "
            "cognitive tests because pollution can affect brain development.\n\n"
            "There is also something called sick building syndrome where the air inside offices and "
            "homes gets polluted from things like paint fumes, cleaning chemicals and dust. So "
            "pollution is not just an outdoor problem."
        )),
        ("What Can Be Done", (
            "On an individual level things like using public transport, not burning garbage, and "
            "planting trees around the house help a little. But I think the real solutions need to "
            "come from the government and industries. Stricter emission standards for vehicles, "
            "switching factories to cleaner energy, and better waste management systems are the "
            "kinds of changes that will actually make a difference at a scale that matters."
        )),
        ("Conclusion", (
            "Air pollution is one of those problems where the effects are slow and invisible at "
            "first but build up over years into serious disease. Living next to the highway has "
            "shown me this in my own family. I think the subject deserves much more attention than "
            "it currently gets in school because it directly affects all of us every single day."
        )),
    ])
    path = os.path.join(OUTPUT_DIR, "Student1_Ravi_Sharma_Assignment3.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}")

# STUDENT 2: Meera Nair — GENUINE ORIGINAL different angle (Expected: CLEAN)
def create_meera():
    doc = Document()
    add_header(doc, "Meera Nair", "10A-034", "Environmental Science",
               "Assignment 3 - Effects of Air Pollution on Human Health", "9 June 2026")
    add_body(doc, [
        ("My Starting Point", (
            "My school is about two kilometres from an industrial area and every time the wind "
            "blows from that direction there is a strange chemical smell in the classrooms. Our "
            "science teacher once told us that smell is sulfur dioxide from the nearby factory. "
            "That conversation is what pushed me to understand what exactly these chemicals do to "
            "our bodies over time."
        )),
        ("Types of Air Pollutants", (
            "There are different categories of air pollutants. Primary pollutants are released "
            "directly into the air like smoke from chimneys or carbon monoxide from bike exhausts. "
            "Secondary pollutants form when primary pollutants react with each other in the "
            "atmosphere. Ozone at ground level is an example of a secondary pollutant. Both types "
            "are harmful but secondary pollutants can sometimes be even harder to predict and control "
            "because they form through atmospheric chemistry that varies with temperature and sunlight."
        )),
        ("Short Term Health Effects", (
            "When someone is exposed to polluted air for even a short period they can experience "
            "burning eyes, scratchy throat, headaches and dizziness. I noticed that during Diwali "
            "when the firecrackers are at their worst, a lot of my friends and I wake up with sore "
            "throats the next morning even though we were indoors most of the night. The air had "
            "crept inside anyway.\n\n"
            "People who already have conditions like asthma or allergies are hit much harder. For "
            "them even a small increase in pollution can mean a full blown attack requiring medication "
            "or hospital treatment."
        )),
        ("Long Term Health Effects", (
            "The longer term picture is quite alarming. Continuous breathing of polluted air over "
            "years gradually damages the tiny air sacs in the lungs called alveoli where oxygen "
            "exchange happens. Once these are damaged they do not heal. This is why people who "
            "worked in coal mines for years almost always develop a disease called black lung.\n\n"
            "Beyond the lungs researchers have now established links between chronic pollution "
            "exposure and cardiovascular disease. The fine particles that enter the bloodstream "
            "cause inflammation in blood vessels which over time leads to plaque buildup and "
            "increases the risk of heart attack and stroke."
        )),
        ("Groups Most at Risk", (
            "Elderly people, young children, pregnant women, and those with pre-existing respiratory "
            "or heart conditions are the most vulnerable. Children are particularly at risk because "
            "their lungs are still developing and they breathe more air relative to their body weight "
            "than adults. Exposure during childhood can permanently reduce lung capacity."
        )),
        ("Conclusion", (
            "Air pollution is a silent health crisis. Its damage accumulates slowly and often by "
            "the time symptoms appear the harm has already been done. Understanding the science "
            "behind how pollutants affect the body is the first step toward demanding better air "
            "quality from our leaders and making smarter personal choices in our daily lives."
        )),
    ])
    path = os.path.join(OUTPUT_DIR, "Student2_Meera_Nair_Assignment3.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}")

# STUDENT 3: Vikram Desai — AI GENERATED (Expected: FLAGGED)
def create_vikram():
    doc = Document()
    add_header(doc, "Vikram Desai", "10A-008", "Environmental Science",
               "Assignment 3 - Effects of Air Pollution on Human Health", "9 June 2026")
    add_body(doc, [
        ("Introduction", (
            "Air pollution constitutes one of the most pressing environmental challenges confronting "
            "contemporary society. It is defined as the presence of harmful substances in the "
            "atmospheric environment in concentrations sufficient to adversely affect human health, "
            "ecosystem integrity, and material properties. The World Health Organization has "
            "classified air pollution as the world largest single environmental health risk, "
            "responsible for approximately seven million premature deaths annually on a global scale."
        )),
        ("Classification of Air Pollutants", (
            "Air pollutants are systematically categorized into primary and secondary classifications. "
            "Primary pollutants are those emitted directly from identifiable sources into the "
            "atmosphere, including particulate matter, sulfur dioxide, nitrogen oxides, carbon "
            "monoxide, and volatile organic compounds. Secondary pollutants, in contrast, are formed "
            "through complex photochemical reactions between primary pollutants and atmospheric "
            "constituents. Tropospheric ozone and secondary particulate matter represent prominent "
            "examples of secondary pollutants with significant implications for human health."
        )),
        ("Respiratory System Impacts", (
            "The respiratory system represents the primary site of interaction between inhaled "
            "pollutants and human physiology. Particulate matter with an aerodynamic diameter of "
            "2.5 micrometers or less, designated PM2.5, penetrates the deepest regions of the "
            "pulmonary system, including the alveolar compartments responsible for gaseous exchange. "
            "Chronic exposure to elevated PM2.5 concentrations is causally associated with the "
            "development of chronic obstructive pulmonary disease, pulmonary fibrosis, and "
            "bronchogenic carcinoma. Furthermore, gaseous pollutants such as sulfur dioxide and "
            "nitrogen dioxide induce airway inflammation and bronchospasm, exacerbating pre-existing "
            "conditions including bronchial asthma."
        )),
        ("Cardiovascular and Systemic Effects", (
            "Epidemiological evidence conclusively demonstrates that air pollution exerts deleterious "
            "effects extending beyond the respiratory system to encompass the cardiovascular system. "
            "Fine particulate matter translocates across the alveolar-capillary membrane into systemic "
            "circulation, where it initiates inflammatory cascades and oxidative stress responses. "
            "These pathophysiological mechanisms contribute substantially to the development of "
            "atherosclerosis, cardiac arrhythmias, myocardial infarction, and ischemic stroke. "
            "Longitudinal cohort studies have consistently demonstrated dose-response relationships "
            "between long-term pollution exposure and cardiovascular mortality."
        )),
        ("Neurological and Developmental Consequences", (
            "Emerging scientific literature has established associations between chronic air pollution "
            "exposure and adverse neurological outcomes. Ultrafine particles and specific gaseous "
            "pollutants possess the capacity to traverse the blood-brain barrier through olfactory "
            "neural pathways, inducing neuroinflammation and oxidative neuronal damage. Pediatric "
            "populations are particularly susceptible, with prenatal and early childhood exposure "
            "demonstrably associated with cognitive developmental delays and reduced academic performance."
        )),
        ("Conclusion", (
            "In conclusion, air pollution represents a multidimensional public health emergency with "
            "ramifications extending across respiratory, cardiovascular, neurological, and "
            "developmental domains. The substantial body of epidemiological and mechanistic evidence "
            "necessitates urgent and comprehensive policy responses at local, national, and "
            "international levels. Sustained commitment to emission reduction, technological "
            "innovation, and health impact surveillance is imperative to mitigate the profound and "
            "pervasive health consequences of atmospheric pollution."
        )),
    ])
    path = os.path.join(OUTPUT_DIR, "Student3_Vikram_Desai_Assignment3.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}")

# STUDENT 4: Deepa Joshi — COPIED FROM RAVI (Expected: FLAGGED)
def create_deepa():
    doc = Document()
    add_header(doc, "Deepa Joshi", "10A-015", "Environmental Science",
               "Assignment 3 - Effects of Air Pollution on Human Health", "9 June 2026")
    add_body(doc, [
        ("Introduction", (
            "I stay near the highway in Pune and the air quality there is really bad. "
            "Some mornings there is a visible haze outside the window and my grandfather always "
            "keeps his inhaler ready before he steps out. That is what made me interested in "
            "this topic. Air pollution is not just a news headline for us, it is something we "
            "face every single day."
        )),
        ("What is Air Pollution", (
            "Air pollution means when harmful substances get mixed into the air we breathe. "
            "These substances can be gases like carbon monoxide and sulfur dioxide, or they can be "
            "tiny solid particles called particulate matter. The main sources are vehicles, factories, "
            "burning of crop residue, and construction dust. In cities like Delhi and Mumbai the "
            "problem is much worse because of the large number of vehicles."
        )),
        ("How it Affects the Human Body", (
            "The most direct effect is on the lungs. When we breathe polluted air the particles go "
            "deep into our lungs and can cause diseases like asthma, bronchitis, and even lung cancer "
            "if the exposure continues for many years. PM2.5 particles are the most dangerous "
            "because they are so small they can enter the lung tissue directly into the "
            "bloodstream. Once in the blood they can reach the heart and brain.\n\n"
            "My grandfather has chronic bronchitis and the doctor told him to avoid going "
            "out on days when the AQI is above 150. I check the AQI app every morning for him. "
            "It made me realise this is not just a general health issue, it affects older "
            "people and children whose lungs are weaker."
        )),
        ("Effects Beyond the Lungs", (
            "What surprised me during my reading was that air pollution does not only affect "
            "breathing. Studies have found links between long term exposure to pollution and heart "
            "disease, stroke, and mental health problems like depression and anxiety. Some "
            "research shows that children growing up in heavily polluted areas score lower on "
            "cognitive tests because pollution can affect brain development.\n\n"
            "There is also something called sick building syndrome where the air inside offices and "
            "homes gets polluted from paint fumes, cleaning chemicals and dust. So "
            "pollution is not just an outdoor problem."
        )),
        ("What Can Be Done", (
            "On an individual level things like using public transport, not burning garbage, and "
            "planting trees around the house help a little. But the real solutions need to "
            "come from the government and industries. Stricter emission standards for vehicles, "
            "switching factories to cleaner energy, and better waste management systems are the "
            "kinds of changes that will make a difference at a scale that matters."
        )),
        ("Conclusion", (
            "Air pollution is one of those problems where the effects are slow and invisible at "
            "first but build up over years into serious illness. I think the subject deserves much "
            "more attention than it gets in school because it directly affects all of us every day."
        )),
    ])
    path = os.path.join(OUTPUT_DIR, "Student4_Deepa_Joshi_Assignment3.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}")

# STUDENT 5: Arjun Kulkarni — PARAPHRASED FROM VIKRAM (Expected: REVIEW)
def create_arjun():
    doc = Document()
    add_header(doc, "Arjun Kulkarni", "10A-042", "Environmental Science",
               "Assignment 3 - Effects of Air Pollution on Human Health", "9 June 2026")
    add_body(doc, [
        ("Introduction", (
            "Dirty air is one of the biggest health problems in today's world. When the "
            "air around us gets filled with toxic gases and tiny particles it becomes harmful to "
            "breathe. The WHO says that bad air kills around 7 million people every year "
            "across the world. It is not a small problem at all."
        )),
        ("Types of Pollutants in the Air", (
            "Scientists divide air pollutants into two main groups. The first group called "
            "primary pollutants are things that go directly into the air from a source like "
            "smoke from factories or exhaust from cars. The second group called secondary pollutants "
            "do not come directly from a source but instead form when primary pollutants mix with "
            "each other in the atmosphere. Ground level ozone is one example of this type."
        )),
        ("Damage to the Lungs", (
            "The lungs take the biggest hit from polluted air. The smallest particles known as "
            "PM2.5 are especially dangerous because they are tiny enough to go all the way into the "
            "deepest parts of the lungs where oxygen enters the blood. Over time this "
            "causes serious lung diseases including long term breathing difficulties and in the "
            "worst cases lung cancer. Gases like sulfur dioxide make breathing even more difficult "
            "and are especially bad for people who already have asthma."
        )),
        ("Heart and Blood Vessel Problems", (
            "Research has clearly shown that polluted air does not just hurt the lungs. The tiny "
            "particles that enter the blood cause swelling and damage inside blood vessels. "
            "This gradually leads to a buildup of fatty deposits in the arteries. Over time "
            "this raises the risk of heart attacks and strokes quite significantly. Large studies "
            "tracking thousands of people over many years confirmed that people in more polluted "
            "areas have higher rates of heart disease."
        )),
        ("Effects on the Brain", (
            "Newer research has found that pollution can even damage the brain. Some of the "
            "smallest particles travel up through the nose directly into the brain "
            "causing inflammation. This is particularly harmful for young children whose brains "
            "are still forming. Kids who grow up in high pollution areas tend to show slower "
            "mental development and perform worse academically."
        )),
        ("Conclusion", (
            "Air pollution is genuinely dangerous and its effects on human health go far deeper "
            "than most people realise. It is not just about coughing and sneezing. It affects "
            "hearts, blood vessels, brains, and even the development of children. Fixing "
            "this needs serious effort from governments, industries, and individuals all working "
            "together toward cleaner air for everyone."
        )),
    ])
    path = os.path.join(OUTPUT_DIR, "Student5_Arjun_Kulkarni_Assignment3.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}")

if __name__ == "__main__":
    print("\n" + "=" * 55)
    print("  Generating 5 Student Assignment Word Files...")
    print("=" * 55)
    create_ravi()
    create_meera()
    create_vikram()
    create_deepa()
    create_arjun()
    print("=" * 55)
    print(f"  All files saved to:\n  {OUTPUT_DIR}")
    print("=" * 55 + "\n")
