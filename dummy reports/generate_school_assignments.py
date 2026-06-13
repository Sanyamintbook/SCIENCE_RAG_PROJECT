"""
generate_school_assignments.py
================================
Generates 5 realistic school student Word assignment files.
Topic: Importance of Water Conservation (Class 10 assignment)

Student Types:
  Kavya  - Researched online, wrote herself  → CLEAN
  Rohit  - Researched online, wrote himself  → CLEAN
  Zara   - Copy pasted from ChatGPT/AI       → FLAGGED (AI)
  Dev    - Copied from Kavya                 → FLAGGED (Copy)
  Nisha  - Rewrote Zara AI text in own words → REVIEW (Paraphrase)

Run: python generate_school_assignments.py
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def make_doc(student_name, roll, subject, topic, date):
    doc = Document()

    # Title
    t = doc.add_heading("ASSIGNMENT", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Info table
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    rows_data = [
        ("Name", student_name),
        ("Roll No.", roll),
        ("Class & Section", "10 - B"),
        ("Subject", subject),
        ("Topic", topic),
    ]
    for i, (k, v) in enumerate(rows_data):
        tbl.rows[i].cells[0].text = k
        tbl.rows[i].cells[1].text = v

    doc.add_paragraph("")
    doc.add_paragraph("Date of Submission: " + date)
    doc.add_paragraph("─" * 58)
    doc.add_paragraph("")
    return doc


def add_section(doc, heading, text):
    if heading:
        h = doc.add_heading(heading, level=2)
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(8)


# ─────────────────────────────────────────────────────────────────────────────
# STUDENT 1 — Kavya Sharma
# Researched on Google, wrote in her OWN words
# Has personal examples, small mistakes, casual tone
# EXPECTED: CLEAN ✅
# ─────────────────────────────────────────────────────────────────────────────
def create_kavya():
    doc = make_doc("Kavya Sharma", "23", "Science", "Importance of Water Conservation", "9 June 2026")

    add_section(doc, "Introduction", (
        "Water is one of the most important things we need to survive. Without water "
        "we cannot live even for few days. I always thought water is unlimited because "
        "we have so many rivers and oceans but when i read about it i was shocked that "
        "only 3 percent of all water on earth is fresh water that we can drink. And most "
        "of that is also locked in glaciers and ice caps so actually very very less water "
        "is available for us to use. This is why saving water is so important."
    ))

    add_section(doc, "Why Water is Getting Scarce", (
        "In our city the tanker comes every two days because tap water supply is not "
        "enough. My mother fills all the buckets and drums when tanker comes. I used "
        "to think this is only our area problem but actually this is happening in many "
        "parts of India. Cities like Chennai and Bengaluru have faced very serious water "
        "crisis in recent years. Farmers in Maharashtra and Rajasthan walk many "
        "kilometers just to get drinking water for their families.\n\n"
        "The reasons water is becoming less are because population is growing so more "
        "people need water, factories and industries use huge amounts of water, and also "
        "because of climate change the rains are not coming properly in many regions. "
        "Also lot of ground water is being pumped out faster than it can refill naturally."
    ))

    add_section(doc, "How We Can Save Water at Home", (
        "There are many small things we can do everyday to save water. I read that "
        "a leaking tap wastes around 20 liters of water per day which is a lot. So "
        "fixing leaks immediately is very important. Also while brushing teeth we should "
        "turn off the tap because water is just flowing for no reason when we are brushing.\n\n"
        "Taking shorter showers instead of bucket baths actually uses more water surprisingly "
        "I thought bucket bath wastes more water but I read that shower in 10 minutes can "
        "use upto 80 liters. Using a bucket and mug is more efficient. Washing clothes "
        "only when there is a full load instead of washing small amounts every day also "
        "saves a lot of water in the washing machine."
    ))

    add_section(doc, "Rainwater Harvesting", (
        "One solution I found very interesting is rainwater harvesting. This means we "
        "collect rainwater from rooftops and store it in tanks or let it go underground "
        "to recharge the water table. In our school they recently build a rainwater "
        "harvesting system on the rooftop and our principal told us it reduces their "
        "water bill by almost 30 percent. This should be made compulsory in all buildings "
        "in my opinion."
    ))

    add_section(doc, "Agriculture and Water Waste", (
        "Agriculture uses about 70 percent of all fresh water in India which is a huge "
        "amount. Traditional flood irrigation method wastes a lot of water because most "
        "of it just evaporates or goes deep into soil without reaching roots. Drip "
        "irrigation is a much better method where water drops directly at plant roots "
        "so almost no wastage. Israel is a country with very little rainfall but they "
        "have become experts in drip irrigation and are food sufficient. India should "
        "also promote this more for farmers."
    ))

    add_section(doc, "Conclusion", (
        "Water conservation is not just government job it is responsibility of every "
        "person. Small actions like closing tap while brushing, fixing leaks, using "
        "bucket instead of pipe for washing cars all add up to big difference. If we "
        "do not start saving water now the future generations will face very serious "
        "shortage. I have started keeping a mug near my basin to collect the water "
        "that runs while waiting for hot water and using it to water our balcony plants. "
        "Everyone should find small ways like this in daily life."
    ))

    path = os.path.join(OUTPUT_DIR, "Student1_Kavya_Sharma_WaterConservation.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}  → Expected: CLEAN")


# ─────────────────────────────────────────────────────────────────────────────
# STUDENT 2 — Rohit Verma
# Researched on Google, wrote in his OWN words — different angle
# Has personal examples, own observations, different structure
# EXPECTED: CLEAN ✅
# ─────────────────────────────────────────────────────────────────────────────
def create_rohit():
    doc = make_doc("Rohit Verma", "31", "Science", "Importance of Water Conservation", "9 June 2026")

    add_section(doc, "What Made Me Think About This", (
        "Last summer we went to my nana's village in Bundelkhand region. The women "
        "there wake up at 4 in the morning just to get water from the hand pump because "
        "by 7 AM there is queue of hundred people. My nani told me this was not the case "
        "20 years back when the pond near the village was full and the well also had water. "
        "Now both are dried up. That trip made me realize that water problem is real and "
        "it is already affecting peoples lives badly."
    ))

    add_section(doc, "The Water Situation in India", (
        "India has about 4 percent of world's fresh water but 18 percent of world's "
        "population. So we are already at a disadvantage. NITI Aayog which is government "
        "planning body said that by 2030 around 40 percent of Indians will not have "
        "access to clean drinking water if current situation continues. This is a very "
        "serious warning.\n\n"
        "The problem is both availability and cleanliness. Many rivers are getting polluted "
        "by factories and sewage. Ganga which is our most sacred river is still very dirty "
        "in many stretches despite government spending thousands of crores on cleaning it. "
        "When rivers get polluted the water cannot be used even for irrigation."
    ))

    add_section(doc, "Different Ways to Conserve", (
        "I researched different conservation methods and found some very creative ones. "
        "In Rajasthan they have a traditional system called Johad which is a community "
        "pond that collects rainwater. These were everywhere 100 years ago but many got "
        "abandoned. Some villages are now reviving them and it has brought water back to "
        "dry areas. A man named Rajendra Singh even got the Stockholm Water Prize for "
        "reviving these traditional water systems. This shows our ancestors already knew "
        "how to manage water well.\n\n"
        "Another thing I learned is about greywater recycling. Greywater is the water "
        "from sinks and washing machines that is not heavily contaminated like toilet water. "
        "This can be cleaned simply and reused for flushing toilets or watering gardens. "
        "In Singapore almost everything is recycled and they have solved their water problem "
        "completely through technology."
    ))

    add_section(doc, "What Schools and Students Can Do", (
        "I think awareness starting from school level is the most important thing. When "
        "children learn good habits young they carry it throughout their life. In our "
        "school there are posters about not wasting water but the taps in the washrooms "
        "are always running and leaking. I think we should have a water monitor system "
        "in school where students track water usage and report wastage. This would make "
        "everyone more conscious.\n\n"
        "At home I have convinced my father to get a low flow showerhead installed. It "
        "reduces water consumption by 40 percent without any difference in experience. "
        "These small changes cost very little but make big impact over time."
    ))

    add_section(doc, "My Opinion", (
        "I think the biggest problem is that water is too cheap in India. Because it "
        "costs so little people do not value it. In countries where water is priced "
        "properly people are much more careful. I am not saying make it unaffordable "
        "for poor people but there should be a system where basic amount is cheap and "
        "excess usage becomes expensive. This would automatically motivate people to save.\n\n"
        "At the end water conservation is about changing mindset. We have to stop thinking "
        "of water as unlimited resource and start treating it like the precious thing it is."
    ))

    path = os.path.join(OUTPUT_DIR, "Student2_Rohit_Verma_WaterConservation.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}  → Expected: CLEAN")


# ─────────────────────────────────────────────────────────────────────────────
# STUDENT 3 — Zara Khan
# Directly submitted ChatGPT/AI generated content
# Perfect formal grammar, no personal voice, academic vocabulary
# EXPECTED: FLAGGED (AI) 🔴
# ─────────────────────────────────────────────────────────────────────────────
def create_zara():
    doc = make_doc("Zara Khan", "07", "Science", "Importance of Water Conservation", "9 June 2026")

    add_section(doc, "Introduction", (
        "Water conservation refers to the strategic management and reduction of water "
        "usage to ensure the sustainable availability of this indispensable natural "
        "resource for current and future generations. As global freshwater reserves "
        "face unprecedented depletion due to population growth, industrialization, "
        "agricultural demand, and climate change, the imperative to conserve water "
        "has become a matter of critical environmental and humanitarian concern. "
        "Approximately 71 percent of the Earth's surface is covered by water; however, "
        "only 2.5 percent constitutes freshwater, and merely a fraction of this is "
        "readily accessible for human consumption and use."
    ))

    add_section(doc, "Causes of Water Scarcity", (
        "Multiple interconnected factors contribute to the escalating global water "
        "crisis. Rapid population growth has substantially increased the demand for "
        "freshwater resources across domestic, agricultural, and industrial sectors. "
        "Concurrently, climate change has disrupted traditional precipitation patterns, "
        "resulting in prolonged droughts in certain regions and flash floods in others, "
        "both of which compromise the reliable availability of freshwater.\n\n"
        "Industrial activities contribute significantly to both water consumption and "
        "contamination. Manufacturing processes, thermal power generation, and chemical "
        "production require vast quantities of water and frequently discharge pollutants "
        "into water bodies, rendering them unsuitable for human or agricultural use. "
        "Furthermore, unsustainable agricultural practices, including excessive groundwater "
        "extraction and inefficient flood irrigation methods, exacerbate water depletion "
        "at an alarming rate."
    ))

    add_section(doc, "Significance of Water Conservation", (
        "The conservation of water resources is of paramount importance for sustaining "
        "both ecological systems and human civilization. Adequate freshwater availability "
        "is fundamental to food security, as agriculture accounts for approximately "
        "70 percent of global freshwater withdrawals. Insufficient water resources "
        "directly threaten agricultural productivity and, consequently, the nutritional "
        "security of populations worldwide.\n\n"
        "Moreover, water conservation plays a vital role in maintaining ecosystem "
        "integrity. Aquatic habitats, wetlands, and riparian ecosystems depend on "
        "consistent water availability to support biodiversity. The degradation of "
        "these ecosystems through water depletion and pollution has far-reaching "
        "consequences for environmental health and human wellbeing."
    ))

    add_section(doc, "Conservation Methods and Strategies", (
        "Effective water conservation necessitates the implementation of comprehensive "
        "strategies across multiple sectors. At the household level, individuals can "
        "significantly reduce water consumption through behavioral modifications such "
        "as fixing plumbing leaks, installing water-efficient fixtures, and adopting "
        "mindful water usage practices. Collectively, these individual actions can "
        "achieve substantial reductions in overall water demand.\n\n"
        "At the infrastructural level, rainwater harvesting systems represent a "
        "particularly effective intervention. By capturing and storing precipitation "
        "for subsequent use, these systems reduce dependence on conventional water "
        "supply networks and contribute to groundwater recharge. Additionally, the "
        "widespread adoption of drip irrigation technology in agriculture can reduce "
        "water usage by up to 50 percent compared to conventional flood irrigation "
        "methods, while simultaneously improving crop yields."
    ))

    add_section(doc, "Conclusion", (
        "In conclusion, water conservation constitutes an indispensable component of "
        "sustainable development and environmental stewardship. The multifaceted "
        "challenges posed by water scarcity demand coordinated responses from "
        "individuals, communities, governments, and international organizations. "
        "Through the systematic implementation of conservation technologies, policy "
        "frameworks, and educational initiatives, it is feasible to ensure equitable "
        "and sustainable access to freshwater resources for present and future "
        "generations. The preservation of this finite and irreplaceable resource "
        "represents a collective responsibility that transcends geographical and "
        "political boundaries."
    ))

    path = os.path.join(OUTPUT_DIR, "Student3_Zara_Khan_WaterConservation.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}  → Expected: FLAGGED (AI generated)")


# ─────────────────────────────────────────────────────────────────────────────
# STUDENT 4 — Dev Patel
# Copied Kavya's assignment with minor word swaps
# EXPECTED: FLAGGED (High copy similarity with Kavya) 🔴
# ─────────────────────────────────────────────────────────────────────────────
def create_dev():
    doc = make_doc("Dev Patel", "12", "Science", "Importance of Water Conservation", "9 June 2026")

    add_section(doc, "Introduction", (
        "Water is one of the most essential things required for survival. Without water "
        "we cannot survive even for a few days. I always believed water was unlimited "
        "because we have so many rivers and oceans but when I read about it I was "
        "shocked to learn that only 3 percent of all water on earth is fresh water "
        "that we can drink. And most of that is also locked in glaciers and ice caps "
        "so very little water is actually available for us. This is why conserving "
        "water is extremely important."
    ))

    add_section(doc, "Why Water is Becoming Scarce", (
        "In our area the water tanker comes every two days because tap water supply "
        "is not adequate. My mother fills all the buckets and drums whenever the tanker "
        "arrives. I used to think this was only a local problem but actually this is "
        "happening in many parts of India. Cities like Chennai and Bengaluru have "
        "experienced very serious water crisis recently. Farmers in Rajasthan and "
        "Maharashtra walk many kilometers just to find drinking water for their families.\n\n"
        "The reasons water is becoming scarce are because population is increasing so "
        "more people require water, industries use huge amounts of water, and also "
        "because of climate change rainfall is not coming properly in many regions. "
        "Additionally a lot of groundwater is being pumped out faster than it can "
        "refill on its own."
    ))

    add_section(doc, "How to Save Water at Home", (
        "There are many small things we can do daily to save water. I read that "
        "a leaking tap wastes approximately 20 liters of water per day which is quite "
        "a lot. So repairing leaks quickly is very important. Also while brushing teeth "
        "we should turn off the tap because water just flows for no reason while we brush.\n\n"
        "Taking bucket baths is more efficient than showers. I thought shower wastes "
        "more water but actually a 10 minute shower can use up to 80 liters. Using a "
        "bucket and mug is more water efficient. Washing clothes only when there is a "
        "full load rather than washing small amounts daily also saves a lot of water."
    ))

    add_section(doc, "Rainwater Harvesting", (
        "One solution I found quite interesting is rainwater harvesting. This means "
        "collecting rainwater from rooftops and storing it in tanks or letting it "
        "recharge the underground water table. In our school they recently built a "
        "rainwater harvesting system and our principal said it reduces their water "
        "bill by nearly 30 percent. This should be made mandatory in all buildings."
    ))

    add_section(doc, "Agriculture Water Usage", (
        "Agriculture uses about 70 percent of all freshwater in India which is a very "
        "large amount. The traditional flood irrigation method wastes a lot of water "
        "because most of it evaporates or goes deep into the soil without reaching "
        "the roots. Drip irrigation is a better method where water drops directly at "
        "plant roots so there is almost no wastage. India should promote this more."
    ))

    add_section(doc, "Conclusion", (
        "Water conservation is not only the government's responsibility it is the duty "
        "of every person. Small habits like turning off the tap while brushing, fixing "
        "leaks, using bucket instead of pipe for washing vehicles all add up to make "
        "a significant difference. If we do not conserve water now future generations "
        "will face serious shortages. I have started collecting water near my basin "
        "and using it to water our plants. Everyone should find small ways in daily life."
    ))

    path = os.path.join(OUTPUT_DIR, "Student4_Dev_Patel_WaterConservation.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}  → Expected: FLAGGED (copied from Kavya)")


# ─────────────────────────────────────────────────────────────────────────────
# STUDENT 5 — Nisha Gupta
# Paraphrased Zara's AI text in simpler student language
# Same meaning as Zara but completely different words
# EXPECTED: REVIEW (paraphrase of Zara) 🟡
# ─────────────────────────────────────────────────────────────────────────────
def create_nisha():
    doc = make_doc("Nisha Gupta", "28", "Science", "Importance of Water Conservation", "9 June 2026")

    add_section(doc, "Introduction", (
        "Saving water means using it carefully and not wasting it so that there is "
        "enough for people now and also for people in the future. The earth has a "
        "lot of water but most of it is salty ocean water. Only about 2.5 percent "
        "is fresh water and only a small part of that is easy to access. So the "
        "actual amount of water we can use is very small. Due to growing population "
        "and pollution this small amount is reducing every year which is why water "
        "conservation has become so necessary."
    ))

    add_section(doc, "Why Fresh Water is Running Out", (
        "There are several reasons why fresh water is becoming difficult to find. "
        "The number of people on earth keeps growing and so the need for water also "
        "increases. Factories and power plants use enormous amounts of water and also "
        "dump chemicals into rivers making them dirty and unusable.\n\n"
        "Weather is also changing because of global warming. Some places get too much "
        "rain causing floods while other places get no rain at all for months. Both "
        "situations make it hard to have steady supply of clean water. Farmers also "
        "use too much groundwater for crops by using old fashioned methods of flooding "
        "fields which wastes most of the water."
    ))

    add_section(doc, "Why Saving Water Matters", (
        "If we do not save water it will affect food production because farming needs "
        "huge amounts of water. Less water means less crops and that means less food "
        "for everyone. Water is also needed by animals plants and entire ecosystems. "
        "When rivers and lakes dry up because of overuse the animals living there die "
        "and the whole balance of nature gets disturbed.\n\n"
        "For humans directly clean drinking water is a basic need. When it becomes "
        "scarce it leads to disease poverty and conflict between communities and even "
        "between countries."
    ))

    add_section(doc, "Ways to Save Water", (
        "At the household level people can save water by being more careful in their "
        "daily habits. Checking for and fixing leaking pipes and taps can save "
        "significant amounts. Using water saving taps and showerheads also helps "
        "reduce consumption without changing lifestyle much.\n\n"
        "Collecting rainwater in tanks is another very effective method. The rain "
        "that falls on rooftops can be directed into storage tanks and used later "
        "for gardening or cleaning. This also helps refill underground water levels. "
        "For farming switching to drip method where water goes directly to plant roots "
        "can cut water use by half compared to traditional flooding methods."
    ))

    add_section(doc, "Conclusion", (
        "To sum up, water conservation is something that all of us must take seriously. "
        "Governments need to make laws and build better systems for managing water. "
        "Farmers need better technology for growing crops with less water. And ordinary "
        "people need to change small habits at home. All these efforts together can "
        "make sure that clean water is available not just for us but for the children "
        "and grandchildren who come after us. Water is not something we can create "
        "so we must protect what we have."
    ))

    path = os.path.join(OUTPUT_DIR, "Student5_Nisha_Gupta_WaterConservation.docx")
    doc.save(path)
    print(f"  Saved: {os.path.basename(path)}  → Expected: REVIEW (paraphrased Zara's AI text)")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("  Generating 5 School Student Assignment Word Files")
    print("  Topic: Importance of Water Conservation (Class 10)")
    print("=" * 60)
    create_kavya()
    create_rohit()
    create_zara()
    create_dev()
    create_nisha()
    print("=" * 60)
    print(f"\n  Files saved to:\n  {OUTPUT_DIR}")
    print("\n  Upload order for testing:")
    print("  1. Student name: Kavya  → Student1_Kavya_Sharma...")
    print("  2. Student name: Rohit  → Student2_Rohit_Verma...")
    print("  3. Student name: Zara   → Student3_Zara_Khan...")
    print("  4. Student name: Dev    → Student4_Dev_Patel...")
    print("  5. Student name: Nisha  → Student5_Nisha_Gupta...")
    print("=" * 60 + "\n")
